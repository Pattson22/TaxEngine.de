"""
Orchestrates one document upload: validate -> store (for provenance) ->
extract. Mirrors `app/eric/submission_service.py`'s role for the ELSTER
scaffold -- the one place the upload bytes, DocumentStorage, and
DocumentExtractionClient meet, so those two stay free of upload-handling
concerns.
"""

from __future__ import annotations

import io
import uuid

import docx

from app.documents.extraction_client import DocumentExtractionClient, DocumentExtractionError
from app.documents.storage import DocumentStorage
from app.schemas.document_extraction import WageCertificateExtraction

MAX_UPLOAD_BYTES = 15 * 1024 * 1024  # 15 MB

_DOCX_CONTENT_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

# Content types accepted from the upload itself -- what extraction_client
# actually sees may differ (a .docx is converted to text/plain below,
# since Claude's document blocks accept PDF/images, not docx bytes).
SUPPORTED_UPLOAD_CONTENT_TYPES = {"application/pdf", "image/png", "image/jpeg", _DOCX_CONTENT_TYPE}


def extract_wage_certificate_from_upload(
    *,
    user_id: uuid.UUID,
    filename: str,
    content_type: str,
    data: bytes,
    storage: DocumentStorage,
    extraction_client: DocumentExtractionClient,
) -> tuple[WageCertificateExtraction, str]:
    """Validates, stores, and extracts one uploaded Lohnsteuerbescheinigung.

    Returns (extraction, storage_key) -- the key is handed back to the
    caller so it can be attached to the WageTaxCertificate the filer
    ultimately saves (see wage_tax_certificates.py's `source_document_key`),
    but nothing here writes to the database itself.
    """
    if content_type not in SUPPORTED_UPLOAD_CONTENT_TYPES:
        raise DocumentExtractionError(
            f"Unsupported file type ({content_type}). Upload a PDF, PNG, JPEG, or Word document."
        )
    if len(data) > MAX_UPLOAD_BYTES:
        raise DocumentExtractionError("That file is too large (max 15 MB).")
    if len(data) == 0:
        raise DocumentExtractionError("That file is empty.")

    safe_filename = filename.replace("/", "_").replace("\\", "_")
    storage_key = f"wage-tax-certificates/{user_id}/{uuid.uuid4()}-{safe_filename}"
    storage.upload(storage_key, data, content_type)

    if content_type == _DOCX_CONTENT_TYPE:
        extraction_data, extraction_media_type = _docx_to_text(data), "text/plain"
    else:
        extraction_data, extraction_media_type = data, content_type

    extraction = extraction_client.extract_wage_certificate(extraction_data, extraction_media_type)
    return extraction, storage_key


def _docx_to_text(data: bytes) -> bytes:
    try:
        document = docx.Document(io.BytesIO(data))
    except Exception as exc:  # python-docx raises plain Exception/PackageNotFoundError for a bad file
        raise DocumentExtractionError("Couldn't open that Word document -- is it a valid .docx file?") from exc

    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    for table in document.tables:
        for row in table.rows:
            text += "\n" + " | ".join(cell.text for cell in row.cells)
    return text.encode("utf-8")
